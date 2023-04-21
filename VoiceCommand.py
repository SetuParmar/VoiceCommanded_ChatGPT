import speech_recognition as sr
import pyttsx3
import openai
import urllib.request
from docx import Document

def takeCommand():
    global answer
    answer = input("Press 1 for text query \n Press 2 for image query ")
    answer = int(answer)
    r = sr.Recognizer()
    with sr.Microphone() as source:
        r.pause_threshold = 1
        r.adjust_for_ambient_noise(source)
        print("Listening...")
        audio = r.listen(source)
    try:
        print("Recognizing...")
        query = r.recognize_google(audio,language= 'en-in')
        #print(query)
        if answer == 1:
            textquery(query)
        else:
            imagequery(query)
    except Exception as e:
        pyttsx3.speak("Chat G P T couldn't recognize what you said, speak once more.")
        print("Chat G P T couldn't recognize what you said, speak once more.")
        return None
def textquery(query):
    document = Document()
    document.add_heading("Welcome to ChatBot", 0)
    openai.api_key = "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"
    document.add_heading(query, 2)
    document.add_heading(" ",2)
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
                {"role": "user", "content": query},
            ]
    )
    result = " "
    for choice in response.choices:
        result += choice.message.content
    pyttsx3.speak(result)
    document.add_paragraph(result)

    document.save('/Users/setuparmar/Documents/AI BOt/Answer.docx')
    print("Your response has been saved to the file !!")

def imagequery(query):
    openai.api_key ='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'
    print(query)
    response = openai.Image.create(
        prompt=query,
        n=1,
        size="1024x1024",
    )

    urk = (response["data"][0]["url"])
    urllib.request.urlretrieve(urk, "/Users/setuparmar/Documents/AI BOt/Image.jpg")
    print("-> Your image has been saved successfully")
takeCommand()